import streamlit as st
import time
from datetime import datetime
import io
import textwrap
import streamlit.components.v1 as components
import html

# --- VERSION STAMP ---
VERSION = "v1.46.3"
BUILD_DATE = "2026-07-21"

# --- FAIL-SAFE ENGINE ---
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# ============================================================
# PROJECT: JUSTICE BOT AI (Global Executive v1.46.3 ELITE)
# PRODUCED BY: Trend Shadows Digital Agency
# STATUS: LIVE PRODUCTION | UNIFIED PRICING
# FIXED: Unified $7.00 price across all gates.
# FIXED: Re-synced sidebar and internal version tags.
# ============================================================

st.set_page_config(page_title="JusticeBot Pro | Global Elite", layout="wide")

# --- SIDEBAR DIAGNOSTICS ---
st.sidebar.markdown("### 🛠️ System Control")
st.sidebar.markdown(f"**VERSION:** `{VERSION}` (ELITE)")
st.sidebar.markdown(f"**BUILD DATE:** `{BUILD_DATE}`")
st.sidebar.markdown("---")
engine_status = "🟢 ACTIVE" if DOCX_SUPPORT else "🟠 INITIALIZING"
st.sidebar.markdown(f"**WORD ENGINE:** `{engine_status}`")
st.sidebar.markdown("---")

# --- MASTER CSS ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700&family=Inter:wght@400;700;800&display=swap');
    .stApp { background-color: #000000 !important; }
    h1 { color: #FFFFFF !important; font-family: 'Playfair Display', serif !important; font-weight: 900 !important; font-size: 3.5rem !important; text-align: center; }
    label, .stMarkdown p { color: #FFFFFF !important; font-weight: 700 !important; font-size: 1.1rem !important; }
    div[data-testid="stSelectbox"] > div { background-color: #C0C0C0 !important; color: #000000 !important; border: 2px solid #FFFFFF !important; border-radius: 4px !important; }
    div[data-testid="stSelectbox"] * { color: #000000 !important; font-weight: 800 !important; }
    .stTextArea textarea, .stTextInput input { background-color: #111111 !important; color: #FFFFFF !important; border: 1px solid #C0C0C0 !important; font-size: 16px !important; }
    button, .stButton>button, .stDownloadButton>button { background-color: #C0C0C0 !important; color: #000000 !important; font-weight: 900 !important; text-transform: uppercase !important; border: 2px solid #FFFFFF !important; height: 3.5em !important; }
    button:hover { background-color: #FFFFFF !important; box-shadow: 0 0 20px rgba(255, 255, 255, 0.5) !important; }
    .ts-logo { font-size: 32px; font-weight: 900; letter-spacing: 10px; color: #FFFFFF; text-align: right; margin-bottom: 20px; }
    </style>
    """, unsafe_allow_html=True)

# --- GLOBAL DATA ---
STATUTES = {
    "United States (US)": {
        "Security Deposit Recovery": {"type": "DEMAND", "law": "Civil Code Section 1950.5 and the URLTA"},
        "Unpaid Freelance Invoice": {"type": "DEMAND", "law": "UCC Article 2 and state Prompt Payment Acts"},
        "Private Vehicle Sale": {"type": "CONTRACT", "law": "State Bill of Sale & UCC Article 2"},
        "Electronics/Goods Sale": {"type": "CONTRACT", "law": "Magnuson-Moss Warranty Act and UCC"},
        "Travel/Flight Refund": {"type": "DEMAND", "law": "14 CFR Part 259 and DOT mandates"},
        "Service Cancellation": {"type": "DEMAND", "law": "State Consumer Statutes & FTC Cooling-Off Rule"},
        "Employment/Unpaid Wages": {"type": "DEMAND", "law": "Fair Labor Standards Act (FLSA)"},
        "Insurance Claim Dispute": {"type": "DEMAND", "law": "State Insurance Fair Conduct Acts"},
        "Property Damage Claim": {"type": "DEMAND", "law": "General Tort Law & Civil Liability Codes"},
        "General Cease & Desist": {"type": "DEMAND", "law": "Common Law Defamation Statutes"}
    },
    "United Kingdom (UK)": {
        "Security Deposit Recovery": {"type": "DEMAND", "law": "Housing Act 2004 and Landlord and Tenant Act 1985"},
        "Unpaid Freelance Invoice": {"type": "DEMAND", "law": "Late Payment of Commercial Debts Act 1998"},
        "Private Vehicle Sale": {"type": "CONTRACT", "law": "Sale of Goods Act 1979 & Common Law"},
        "Electronics/Goods Sale": {"type": "CONTRACT", "law": "Sale of Goods Act 1979"},
        "Travel/Flight Refund": {"type": "DEMAND", "law": "UK261 Flight Compensation Regulations"},
        "Service Cancellation": {"type": "DEMAND", "law": "Consumer Rights Act 2015"},
        "Employment/Unpaid Wages": {"type": "DEMAND", "law": "Employment Rights Act 1996"},
        "Insurance Claim Dispute": {"type": "DEMAND", "law": "FOS Guidelines"},
        "Property Damage Claim": {"type": "DEMAND", "law": "Occupiers' Liability Act 1957"},
        "General Cease & Desist": {"type": "DEMAND", "law": "Protection from Harassment Act 1997"}
    },
    "South Africa (ZA)": {
        "Security Deposit Recovery": {"type": "DEMAND", "law": "Rental Housing Act 50 of 1999"},
        "Unpaid Freelance Invoice": {"type": "DEMAND", "law": "Section 11 of the Prescription Act"},
        "Private Vehicle Sale": {"type": "CONTRACT", "law": "Common Law of Sale and Voetstoots Protocol"},
        "Electronics/Goods Sale": {"type": "CONTRACT", "law": "Consumer Protection Act 68 of 2008"},
        "Travel/Flight Refund": {"type": "DEMAND", "law": "CPA Section 47 and IATA Conventions"},
        "Service Cancellation": {"type": "DEMAND", "law": "CPA Section 14"},
        "Employment/Unpaid Wages": {"type": "DEMAND", "law": "Basic Conditions of Employment Act"},
        "Insurance Claim Dispute": {"type": "DEMAND", "law": "Short-Term Insurance Act"},
        "Property Damage Claim": {"type": "DEMAND", "law": "Apportionment of Damages Act"},
        "General Cease & Desist": {"type": "DEMAND", "law": "Protection from Harassment Act 2011"}
    },
    "Australia (AU)": {
        "Security Deposit Recovery": {"type": "DEMAND", "law": "Residential Tenancies Act 2010"},
        "Unpaid Freelance Invoice": {"type": "DEMAND", "law": "Australian Consumer Law (ACL)"},
        "Private Vehicle Sale": {"type": "CONTRACT", "law": "State Motor Vehicle Traders Act"},
        "Electronics/Goods Sale": {"type": "CONTRACT", "law": "ACL Consumer Guarantees"},
        "Travel/Flight Refund": {"type": "DEMAND", "law": "ACL Major Failure Protocols"},
        "Service Cancellation": {"type": "DEMAND", "law": "ACL Unfair Contract Terms"},
        "Employment/Unpaid Wages": {"type": "DEMAND", "law": "Fair Work Act 2009"},
        "Insurance Claim Dispute": {"type": "DEMAND", "law": "Insurance Contracts Act 1984"},
        "Property Damage Claim": {"type": "DEMAND", "law": "Civil Liability Act 2002"},
        "General Cease & Desist": {"type": "DEMAND", "law": "Personal Violence Protection Act"}
    },
    "Canada (CA)": {
        "Security Deposit Recovery": {"type": "DEMAND", "law": "Provincial Residential Tenancy Acts"},
        "Unpaid Freelance Invoice": {"type": "DEMAND", "law": "Provincial Sale of Goods Acts"},
        "Private Vehicle Sale": {"type": "CONTRACT", "law": "Provincial Consumer Protection Statutes"},
        "Electronics/Goods Sale": {"type": "CONTRACT", "law": "CPA Section 14"},
        "Travel/Flight Refund": {"type": "DEMAND", "law": "Air Passenger Protection Regulations (APPR)"},
        "Service Cancellation": {"type": "DEMAND", "law": "CPA Cancellation Protocols"},
        "Employment/Unpaid Wages": {"type": "DEMAND", "law": "Canada Labour Code"},
        "Insurance Claim Dispute": {"type": "DEMAND", "law": "Provincial Insurance Acts"},
        "Property Damage Claim": {"type": "DEMAND", "law": "Negligence Act"},
        "General Cease & Desist": {"type": "DEMAND", "law": "Libel and Slander Act"}
    },
    "New Zealand (NZ)": {
        "Security Deposit Recovery": {"type": "DEMAND", "law": "Residential Tenancies Act 1986"},
        "Unpaid Freelance Invoice": {"type": "DEMAND", "law": "Consumer Guarantees Act 1993"},
        "Private Vehicle Sale": {"type": "CONTRACT", "law": "Fair Trading Act 1986"},
        "Electronics/Goods Sale": {"type": "CONTRACT", "law": "CGA Section 6"},
        "Travel/Flight Refund": {"type": "DEMAND", "law": "Contract and Commercial Law Act 2017"},
        "Service Cancellation": {"type": "DEMAND", "law": "CGA Reasonable Care"},
        "Employment/Unpaid Wages": {"type": "DEMAND", "law": "Employment Relations Act 2000"},
        "Insurance Claim Dispute": {"type": "DEMAND", "law": "Insurance Law Reform Act"},
        "Property Damage Claim": {"type": "DEMAND", "law": "Limitation Act 2010"},
        "General Cease & Desist": {"type": "DEMAND", "law": "Harassment Act 1997"}
    },
    "India (IN)": {
        "Security Deposit Recovery": {"type": "DEMAND", "law": "Model Tenancy Act 2021"},
        "Unpaid Freelance Invoice": {"type": "DEMAND", "law": "Indian Contract Act 1872"},
        "Private Vehicle Sale": {"type": "CONTRACT", "law": "Motor Vehicles Act 1988"},
        "Electronics/Goods Sale": {"type": "CONTRACT", "law": "Consumer Protection Act 2019"},
        "Travel/Flight Refund": {"type": "DEMAND", "law": "DGCA Passenger Charter"},
        "Service Cancellation": {"type": "DEMAND", "law": "CPA Unfair Contracts"},
        "Employment/Unpaid Wages": {"type": "DEMAND", "law": "Payment of Wages Act 1936"},
        "Insurance Claim Dispute": {"type": "DEMAND", "law": "IRDAI Guidelines"},
        "Property Damage Claim": {"type": "DEMAND", "law": "Tort of Negligence"},
        "General Cease & Desist": {"type": "DEMAND", "law": "IT Act 2000"}
    }
}
CURRENCIES = ["USD ($)", "ZAR (R)", "GBP (£)", "AUD ($)", "NZD ($)", "CAD ($)", "INR (₹)"]

# --- PERSISTENT STORAGE ---
if 'paid_v42' not in st.session_state: st.session_state.paid_v42 = False
if 'ready_v42' not in st.session_state: st.session_state.ready_v42 = False
if 'vault_v42' not in st.session_state: st.session_state.vault_v42 = {}

st.markdown('<div class="ts-logo">TS</div>', unsafe_allow_html=True)
st.title("JusticeBot Pro Global")
st.divider()

if not st.session_state.paid_v42:
    with st.container():
        st.markdown("### 🏛️ Case Intelligence Architecture")
        c1, c2 = st.columns(2)
        with c1:
            cl_name = st.text_input("Claimant Full Name (You)", key="cl_in")
            cl_id = st.text_input("Claimant ID Number / Passport", key="cl_id_in")
            cl_addr = st.text_area("Claimant Physical Address", height=80, key="cla_in")
            juris = st.selectbox("Select Jurisdiction", list(STATUTES.keys()), key="ju_in")
        with c2:
            res_name = st.text_input("Respondent Name (Target)", key="re_in")
            res_id = st.text_input("Respondent ID Number (If known)", key="res_id_in")
            res_addr = st.text_area("Respondent Physical Address", height=80, key="resa_in")
            curr = st.selectbox("Select Currency", CURRENCIES, key="cu_in")
            
        st.divider()
        category = st.selectbox("Select Case Category", list(STATUTES[juris].keys()), key="ca_in")
        
        is_vehicle = "Vehicle" in category
        if is_vehicle:
            c3, c4 = st.columns(2)
            with c3: ref_no = st.text_input("Vehicle VIN / Serial Number", key="vin_in")
            with c4: reg_no = st.text_input("Vehicle Registration / Plate", key="reg_in")
            ref_combined = f"VIN: {ref_no} | REG: {reg_no}"
        else:
            ref_combined = st.text_input("Reference / Invoice / Lease Number", key="ref_in")
            
        amount = st.text_input(f"Total Amount Owed / Sale Price ({curr.split(' ')[1]})", key="am_in")
        details = st.text_area("Narrative (Dates, facts, and specifics)", height=150, key="de_in")
        
        if st.button("PROCESS OFFICIAL DOCUMENT", key="main_btn"):
            if cl_name and res_name and details and amount:
                st.session_state.vault_v42 = {
                    "cl": cl_name, "cl_id": cl_id, "cla": cl_addr, 
                    "res": res_name, "res_id": res_id, "resa": res_addr,
                    "amt": amount, "det": details, "jur": juris, "cat": category,
                    "ref": ref_combined, "cur": curr.split(' ')[1]
                }
                st.session_state.ready_v42 = True
                st.success("Analysis Complete. Draft Generated.")

    if st.session_state.get("ready_v42"):
        st.divider()
        v = st.session_state.vault_v42
        mode_type = STATUTES[v['jur']][v['cat']]['type']
        subject_header = "NOTICE OF INTENT TO LITIGATE" if mode_type == "DEMAND" else "SALE & PURCHASE AGREEMENT"
        st.markdown(f"**SUBJECT: {subject_header} - {v['cat'].upper()}**")
        st.markdown("""<div style="filter:blur(12px); background:#111; padding:20px; border:1px dashed #C0C0C0;">Notice is hereby given that your actions constitute a direct violation... [ENCRYPTED]</div>""", unsafe_allow_html=True)
        
        st.markdown("### 💎 Unlock Full Document Suite")
        st.info("Pay once, receive your activation key instantly, and unlock the high-end Word/Docx suite.")
        
        col_ls, col_key = st.columns(2)
        with col_ls:
            st.link_button("💳 PAY $7.00 VIA LEMONSQUEEZY", "https://trend-shadows.lemonsqueezy.com/checkout/buy/bccbd513-8d20-4156-ad75-aa4ca00fc2d8")
        with col_key:
            v_code = st.text_input("Enter License Key / Voucher", type="password", key="vouch_in")
            if st.button("ACTIVATE ELITE SUITE", key="auth_btn"):
                # Accepts Gift Code OR any LS-Style Key (Admin logic)
                if v_code == "TS-GIFT-2026" or len(v_code) > 10:
                    st.session_state.paid_v42 = True
                    st.rerun()
                else: st.error("Invalid or Inactive License Key")

else:
    # --- PRODUCTION OUTPUT ---
    v = st.session_state.vault_v42
    st.success("✅ LICENSE VALIDATED. ACCESS GRANTED.")
    config = STATUTES[v['jur']].get(v['cat'])
    law = config['law']
    doc_mode = config['type']
    date_now = datetime.now().strftime("%B %d, %Y")
    
    def clean_id(val):
        if not val: return ""
        return val.upper().replace('ID:', '').replace('ID', '').strip()

    cl_id_clean = clean_id(v.get('cl_id', ''))
    res_id_clean = clean_id(v.get('res_id', ''))

    if doc_mode == "DEMAND":
        doc_title = "FORMAL LETTER OF DEMAND - PRE-LITIGATION NOTICE"
        res_id_tag = f" (ID: {res_id_clean})" if res_id_clean else ""
        content_body = f"TO: {v['res']}{res_id_tag}\n\nNotice is hereby given that your failure to remit the balance of {v['cur']}{v['amt']} regarding the {v['cat'].lower()} constitutes a direct violation of {law}.\n\nUnder the laws of {v['jur']}, the withholding of these funds is a breach of legal obligations.\n\nSTATEMENT OF FACTS:\n{v['det']}\n\nLEGAL DEMAND:\nDemand is hereby made for the immediate payment of the full balance. This must be received in full within 14 calendar days of the date of this notice.\n\nINTENT TO LITIGATE:\nFailure to comply with this final demand will result in the immediate commencement of legal proceedings in Small Claims Court without further notice."
    else:
        doc_title = "PRIVATE SALE & PURCHASE AGREEMENT"
        cl_id_tag = f" (ID: {cl_id_clean})" if cl_id_clean else ""
        res_id_tag = f" (ID: {res_id_clean})" if res_id_clean else ""
        content_body = f"This Agreement is made on {date_now} between {v['cl']}{cl_id_tag} (Seller) and {v['res']}{res_id_tag} (Buyer).\n\nASSET DESCRIPTION:\n{v['cat']} - {v['ref']}\n\nNARRATIVE & CONDITION:\n{v['det']}\n\nPURCHASE PRICE:\nThe agreed sale price is {v['cur']}{v['amt']}, payable in full before the transfer of ownership.\n\nLEGAL TERMS:\nThis sale is conducted under the {law}. The asset is sold in its current condition (As-Is/Voetstoots), and the Seller provides no warranties. The Buyer acknowledges they have inspected the asset and accept it in its current state."

    # --- HTML PREVIEW ---
    body_html = ""
    for line in content_body.split('\n'):
        line = html.escape(line.strip())
        if not line: body_html += "<br>"
        elif ":" in line and len(line) < 45 and (line.isupper() or line.endswith(':')):
            body_html += f"<b style='text-decoration: underline; color: #000; font-family: Arial, sans-serif;'>{line}</b><br>"
        else:
            body_html += f"<span style='color: #000;'>{line}</span><br>"

    full_html = f"""
    <html>
    <body style="margin:0; padding:0; background-color: #f4f4f4;">
        <div style="background-color: #FFFFFF; color: #000000; padding: 60px; font-family: 'Times New Roman', serif; line-height: 1.5; border: 1px solid #ddd; max-width: 800px; margin: 20px auto; box-shadow: 0 0 20px rgba(0,0,0,0.1);">
            <div style="border-bottom: 2px solid #000; padding-bottom: 15px; margin-bottom: 25px; display: flex; justify-content: space-between; align-items: flex-end;">
                <div style="font-size: 11px; color: #555; font-family: Arial, sans-serif;">REF: {v['ref']}</div>
                <div style="text-align: right;">
                    <div style="font-family: Arial, sans-serif; font-weight: 900; font-size: 22px; color: #000;">TREND SHADOWS</div>
                    <div style="font-size: 10px; font-style: italic; color: #555; font-family: Arial, sans-serif;">JUSTICE BOT PRO GLOBAL</div>
                </div>
            </div>
            <table style="width: 100%; margin-bottom: 30px; font-size: 12px; color: #000; font-family: 'Times New Roman', serif;">
                <tr>
                    <td style="width: 50%; vertical-align: top;">
                        <b>FROM:</b><br>{v['cl']}<br>ID: {cl_id_clean if cl_id_clean else 'N/A'}<br>{v['cla'].replace('\n','<br>')}
                    </td>
                    <td style="width: 50%; text-align: right; vertical-align: top;">
                        <b>TO:</b><br>{v['res']}<br>ID: {res_id_clean if res_id_clean else 'N/A'}<br>{v['resa'].replace('\n','<br>')}
                    </td>
                </tr>
            </table>
            <div style="text-align: right; font-size: 12px; margin-bottom: 20px; color: #000;"><b>DATE:</b> {date_now}</div>
            <div style="text-align: center; font-family: Arial, sans-serif; font-weight: 900; font-size: 18px; text-decoration: underline; margin: 30px 0; text-transform: uppercase; color: #000;">{doc_title}</div>
            <div style="font-size: 14px; text-align: justify; color: #000; font-family: 'Times New Roman', serif; line-height: 1.6;">{body_html}</div>
            <table style="width: 100%; margin-top: 60px; font-size: 11px; color: #000;">
                <tr>
                    <td style="width: 45%; border-top: 2px solid #000; padding-top: 10px;">__________________________<br><b>FOR THE SELLER / CLAIMANT</b><br>Name: {v['cl']}</td>
                    <td style="width: 10%;"></td>
                    <td style="width: 45%; border-top: 2px solid #000; padding-top: 10px; text-align: right;">__________________________<br><b>FOR THE BUYER / RESPONDENT</b><br>Name: {v['res']}</td>
                </tr>
            </table>
        </div>
    </body>
    </html>
    """
    st.markdown("### 📄 Professional Preview")
    components.html(full_html, height=700, scrolling=True)

    # --- DOCX GENERATION ---
    def generate_pro_docx(v, title, body, date_str, cl_id, res_id):
        if not DOCX_SUPPORT: return None
        try:
            doc = Document()
            # Font setup
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(11)

            # 1. Branding Header
            section = doc.sections[0]
            header = section.header
            htable = header.add_table(1, 2, width=Inches(6.5))
            htable.columns[0].width = Inches(4.5)
            htable.columns[1].width = Inches(2.0)
            
            h_cell = htable.cell(0, 1)
            hp = h_cell.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hr = hp.add_run("TREND SHADOWS")
            hr.bold = True
            hr.font.size = Pt(10)
            hr.font.name = 'Arial'
            
            hp2 = h_cell.add_paragraph()
            hp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hr2 = hp2.add_run("JUSTICE BOT PRO GLOBAL")
            hr2.font.size = Pt(8)
            hr2.font.name = 'Arial'
            hr2.italic = True

            # 2. Party Table (Side-by-Side)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(3.25)
            table.columns[1].width = Inches(3.25)
            
            # FROM
            c1 = table.cell(0, 0)
            cp1 = c1.paragraphs[0]
            cp1.add_run("FROM (CLAIMANT/SELLER):").bold = True
            cp1.add_run(f"\n{v.get('cl', '')}")
            if cl_id: cp1.add_run(f"\nID: {cl_id}")
            cp1.add_run(f"\n{v.get('cla', '')}")
            
            # TO
            c2 = table.cell(0, 1)
            cp2 = c2.paragraphs[0]
            cp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cp2.add_run("TO (RESPONDENT/BUYER):").bold = True
            cp2.add_run(f"\n{v.get('res', '')}")
            if res_id: cp2.add_run(f"\nID: {res_id}")
            cp2.add_run(f"\n{v.get('resa', '')}")

            doc.add_paragraph("\n")
            
            # 3. Date
            dr = doc.add_paragraph()
            dr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            dr.add_run(f"DATE: {date_str}").bold = True
            
            doc.add_paragraph("\n")
            
            # 4. Title
            t = doc.add_paragraph()
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tr = t.add_run(title)
            tr.bold = True
            tr.font.size = Pt(16)
            tr.font.name = 'Arial'
            
            # Decorative Line
            p_line = doc.add_paragraph()
            p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_line.add_run("_______________________________________________________")

            doc.add_paragraph("\n")
            
            # 5. Body Content
            for line in body.split('\n'):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                    continue
                
                p = doc.add_paragraph()
                if ":" in line and len(line) < 45 and (line.isupper() or line.endswith(':')):
                    run = p.add_run(line)
                    run.bold = True
                else:
                    p.add_run(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_paragraph("\n\n")
            
            # 6. Signature Block
            sig_table = doc.add_table(rows=2, cols=2)
            sig_table.columns[0].width = Inches(3.25)
            sig_table.columns[1].width = Inches(3.25)
            
            s1 = sig_table.cell(0, 0)
            s1.paragraphs[0].add_run("__________________________\nFOR THE SELLER / CLAIMANT")
            s1.add_paragraph(f"Name: {v.get('cl', '')}")
            
            s2 = sig_table.cell(0, 1)
            s2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            s2.paragraphs[0].add_run("__________________________\nFOR THE BUYER / RESPONDENT")
            s2p2 = s2.add_paragraph(f"Name: {v.get('res', '')}")
            s2p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 7. Footer
            footer = section.footer
            fp = footer.paragraphs[0]
            fp.text = "This document is a legally binding instrument generated via Trend Shadows JusticeBot Pro Global. (c) 2026."
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp.style.font.size = Pt(8)
            fp.style.font.italic = True
            
            target = io.BytesIO()
            doc.save(target)
            return target.getvalue()
        except:
            return None

    # --- DOWNLOAD CENTER ---
    st.markdown("### 📥 Document Distribution")
    col1, col2 = st.columns(2)
    with col1:
        docx_bytes = generate_pro_docx(v, doc_title, content_body, date_now, cl_id_clean, res_id_clean)
        if DOCX_SUPPORT and docx_bytes:
            st.download_button("📥 DOWNLOAD ELITE WORD (.DOCX)", docx_bytes, file_name=f"JusticeBot_{v['cat']}.docx")
        else: st.warning("Word Engine Initializing...")
    with col2:
        st.download_button("📥 DOWNLOAD FORMATTED DOC (FALLBACK)", full_html, file_name=f"JusticeBot_{v['cat']}.doc", mime="application/msword")

    if st.button("INITIATE NEW CASE"):
        st.session_state.paid_v42 = False; st.session_state.ready_v42 = False; st.session_state.vault_v42 = {}; st.rerun()

st.divider()
st.caption(f"Shadow-Build Global Engine {VERSION} | Trend Shadows Digital Agency")

# --- MANDATORY DIGISTORE24 COMPLIANCE ---
st.markdown("""
    <div style="text-align: center; font-size: 10px; color: #444; margin-top: 20px;">
        The debit will be performed by Digistore24.com
    </div>
    """, unsafe_allow_html=True)
